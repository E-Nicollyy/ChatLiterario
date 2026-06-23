import os
import re
import google.generativeai as genai
from docx import Document
from PIL import Image
import gradio as gr
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURAÇÃO SEGURA DA API KEY ---
# Busca a chave nas variáveis de ambiente do Hugging Face. 
# Se não encontrar, usa a sua chave antiga como alternativa.
api_key = os.environ.get("GEMINI_API_KEY")

if not api_key:
    raise ValueError("GEMINI_API_KEY não configurada.")

genai.configure(api_key=api_key)
model = genai.GenerativeModel("gemini-2.5-flash")
# --- FUNÇÕES AUXILIARES DE ESTILIZAÇÃO E FORMATO ---

def format_as_card(text, title="Resultado"):
    """Envolve o retorno da inteligência artificial dentro de um card HTML elegante."""
    return f"""
    <div style="background-color: #ffffff; border-radius: 12px; padding: 25px; 
                box-shadow: 0 4px 6px rgba(0,0,0,0.1); border-left: 6px solid #4A90E2; 
                margin: 10px 0; color: #2c3e50; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;">
        <h3 style="margin-top: 0; color: #4A90E2; border-bottom: 1px solid #eee; padding-bottom: 10px;">{title}</h3>
        <div style="line-height: 1.6;">
            {text.replace('\n', '<br>')}
        </div>
    </div>
    """

def save_word_document(content, genre, mode="criacao"):
    """Estrutura e salva os arquivos Word, retornando o caminho do arquivo gerado."""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    filename = f"{mode}_{genre.replace(' ', '_')}.docx"
    
    if mode == "criacao":
        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t = p_titulo.add_run("Texto Literário Original\n")
        run_t.bold = True
        run_t.font.size = Pt(16)
    else:
        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t = p_titulo.add_run(f"Revisão Editorial: {genre}\n")
        run_t.bold = True

    for para in content.split('\n'):
        if para.strip():
            p = doc.add_paragraph(para)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
    doc.save(filename)
    return filename

# --- FUNÇÕES DE PROCESSAMENTO (LÓGICA ORIGINAL) ---

def menu_indicacoes(genero, favorito, imagem, texto_extra):
    prompt = f"""
    Você é um especialista literário.
    Gosto do usuário:
    - Gênero literário favorito: {genero}
    - Livro favorito: {favorito}
    
    Informações adicionais:
    - Comentário extra: "{texto_extra}"
    Tarefa: 
    1. Se houver imagem, descreva detalhadamente e faça a relação literária. Se houver apenas texto, faça a análise do livro/texto informado.
    2. Indique livros parecidos com os gostos do usuário ({genero} e {favorito}).
    3. Comente detalhadamente sobre suas indicações.
    """
    if imagem:
        response = model.generate_content([imagem, prompt])
    else:
        response = model.generate_content(prompt)
    
    return format_as_card(response.text, "📚 Indicações e Análise")

def menu_criacao(genero, imagem, texto_ideia):
    prompt = f"""
Você é um escritor, romancista, poeta e roteirista profissional.
Crie um texto literário original seguindo as informações abaixo.
Gênero: {genero}
Ideia do usuário: {texto_ideia}
Regras:
1. Utilize a estrutura correta para esse gênero.
2. Não explique o que está fazendo.
3. Entregue exatamente com as tags [TITULO], [GENERO] e [TEXTO_COMPLETO].
Estrutura da resposta:
[TITULO]Título[/TITULO]
[GENERO]{genero}[/GENERO]
[TEXTO_COMPLETO]O texto aqui[/TEXTO_COMPLETO]
"""
    if imagem:
        response = model.generate_content([imagem, prompt])
    else:
        response = model.generate_content(prompt)
    
    res_text = response.text
    match = re.search(r"\[TEXTO_COMPLETO\](.*?)\[/TEXTO_COMPLETO\]", res_text, re.DOTALL)
    
    file_path = None
    if match:
        texto_final = match.group(1).strip()
        file_path = save_word_document(texto_final, genero, "criacao")
    
    return format_as_card(res_text, "✍️ Obra Gerada"), file_path

def menu_interpretacao(imagem, texto):
    prompt = f"""
    Você é um crítico literário e analista textual sênior.
    Texto/Descrição: "{texto}"
    Tarefa: INTERPRETE DE UM JEITO DETALHADO. Divida estritamente em:
    1. PARTES SEPARADAS
    2. QUAIS IMPACTOS
    3. DESTAQUE FRASES MAIS IMPACTANTES
    4. O QUE O TEXTO QUER DIZER
    """
    if imagem:
        response = model.generate_content([imagem, prompt])
    else:
        response = model.generate_content(prompt)
        
    return format_as_card(response.text, "🔍 Interpretação Detalhada")

def menu_correcao(genero, imagem, texto):
    prompt = f"""
    Você é um revisor especialista em literatura do gênero: {genero}.
    Texto: "{texto}"
    Retorne sua resposta estritamente seguindo a estrutura abaixo:
    **Apontamentos de Correção**:
    (observações)
    **Texto Sugerido**:
    [TEXTO_FINAL]
    (Apenas o texto corrigido aqui)
    [/TEXTO_FINAL]
    """
    if imagem:
        response = model.generate_content([imagem, prompt])
    else:
        response = model.generate_content(prompt)
        
    res_text = response.text
    match = re.search(r"\[TEXTO_FINAL\](.*?)\[/TEXTO_FINAL\]", res_text, re.DOTALL)
    
    file_path = None
    if match:
        texto_final = match.group(1).strip()
        file_path = save_word_document(texto_final, genero, "revisao")
        
    return format_as_card(res_text, "📝 Correção e Revisão"), file_path

# --- INTERFACE VISUAL (GRADIO EM TEMA MONOCHROME) ---

with gr.Blocks(theme=gr.Theme.from_hub("Luminia/llm-trainer")) as app:
    gr.Markdown("# 📖 Chat Literário Inteligente")
    gr.Markdown("Seja bem-vindo ao seu assistente literário pessoal. Navegue pelas abas abaixo:")
    
    with gr.Tabs():
        # TAB 1: INDICAÇÕES
        with gr.TabItem("📚 Indicações"):
            with gr.Row():
                with gr.Column():
                    in_gen = gr.Textbox(label="Qual gênero literário você mais gosta de ler?")
                    in_fav = gr.Textbox(label="Qual é o seu livro favorito?")
                    in_img = gr.Image(type="pil", label="Foto ou arquivo do livro (Opcional)")
                    in_txt = gr.Textbox(label="Comente algo breve sobre o livro/imagem")
                    btn_ind = gr.Button("Obter Indicações e Análise", variant="primary")
                with gr.Column():
                    out_ind = gr.HTML(label="Resultado da IA")
            btn_ind.click(menu_indicacoes, [in_gen, in_fav, in_img, in_txt], out_ind)

        # TAB 2: CRIAÇÃO LITERÁRIA
        with gr.TabItem("✍️ Criação"):
            with gr.Row():
                with gr.Column():
                    cr_gen = gr.Textbox(label="Qual gênero deseja criar? (Ex: cartas, crônicas, poesias...)")
                    cr_img = gr.Image(type="pil", label="Foto ou arquivo como ideia inicial (Opcional)")
                    cr_txt = gr.Textbox(label="Digite ou cole aqui o texto/resumo da sua ideia inicial", lines=3)
                    btn_cr = gr.Button("Gerar Obra Literária", variant="primary")
                with gr.Column():
                    out_cr_txt = gr.HTML(label="Obra Gerada pela IA")
                    out_cr_file = gr.File(label="Baixar Arquivo Word Gerado (.docx)")
            btn_cr.click(menu_criacao, [cr_gen, cr_img, cr_txt], [out_cr_txt, out_cr_file])

        # TAB 3: INTERPRETAÇÃO TEXTUAL
        with gr.TabItem("🔍 Interpretação"):
            with gr.Row():
                with gr.Column():
                    it_img = gr.Image(type="pil", label="Enviar imagem para interpretação (Opcional)")
                    it_txt = gr.Textbox(label="Digite ou cole aqui o texto que deseja interpretar", lines=5)
                    btn_it = gr.Button("Iniciar Interpretação", variant="primary")
                with gr.Column():
                    out_it = gr.HTML(label="Análise da IA")
            btn_it.click(menu_interpretacao, [it_img, it_txt], out_it)

        # TAB 4: CORREÇÕES
        with gr.TabItem("📝 Correção"):
            with gr.Row():
                with gr.Column():
                    co_gen = gr.Textbox(label="Qual é o gênero do texto para correção?")
                    co_img = gr.Image(type="pil", label="Enviar imagem/arquivo do texto (Opcional)")
                    co_txt = gr.Textbox(label="Digite ou cole aqui o texto que deseja corrigir", lines=5)
                    btn_co = gr.Button("Corrigir e Formatar", variant="primary")
                with gr.Column():
                    out_co_txt = gr.HTML(label="Texto Revisado")
                    out_co_file = gr.File(label="Baixar Texto Corrigido (.docx)")
            btn_co.click(menu_correcao, [co_gen, co_img, co_txt], [out_co_txt, out_co_file])

# Execução padrão exigida para servidores de produção
if __name__ == "__main__":
    app.launch()